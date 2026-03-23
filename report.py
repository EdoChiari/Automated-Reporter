import psycopg2
import anthropic
from dotenv import load_dotenv
from docx import Document
from datetime import datetime
import os
import io
import smtplib
import markdown
from html2docx import html2docx
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders

load_dotenv()

# Read KPIs from file
with open("kpis.txt", "r") as f:
    kpis = [line.strip() for line in f.readlines() if line.strip()]

print(f"Loaded {len(kpis)} KPIs")

# Connect to the database
conn = psycopg2.connect(
    host=os.getenv("DB_HOST"),
    port=os.getenv("DB_PORT"),
    dbname=os.getenv("DB_NAME"),
    user=os.getenv("DB_USER"),
    password=os.getenv("DB_PASSWORD")
)
cursor = conn.cursor()
print("Connected to database!")

# Read schema
cursor.execute("""
    SELECT table_name, column_name, data_type
    FROM information_schema.columns
    WHERE table_schema = 'public'
    ORDER BY table_name, ordinal_position
""")

schema = {}
for table, column, dtype in cursor.fetchall():
    if table not in schema:
        schema[table] = []
    schema[table].append(f"{column} ({dtype})")

schema_text = ""
for table, columns in schema.items():
    schema_text += f"\nTable: {table}\n"
    schema_text += "\n".join([f"  - {col}" for col in columns])
    schema_text += "\n"

print("Schema loaded!")

client = anthropic.Anthropic()
kpi_results = []

# Process each KPI
for kpi in kpis:
    print(f"\nProcessing: {kpi}")

    message = client.messages.create(
        model="claude-sonnet-4-6",
        max_tokens=1024,
        messages=[
            {"role": "user", "content": f"""
            You are a SQL expert. Given this database schema:
            {schema_text}
            
            Write a PostgreSQL query to calculate: {kpi}
            
            Return ONLY the SQL query, nothing else. No markdown, no explanation.
            """}
        ]
    )

    sql_query = message.content[0].text.strip()
    if sql_query.startswith("```"):
        sql_query = sql_query.split("\n", 1)[1]
    if sql_query.endswith("```"):
        sql_query = sql_query.rsplit("```", 1)[0].strip()

    try:
        cursor = conn.cursor()
        cursor.execute(sql_query)
        results = cursor.fetchall()
        column_names = [desc[0] for desc in cursor.description]

        results_text = ", ".join(column_names) + "\n"
        for row in results:
            results_text += str(row) + "\n"

        kpi_results.append({
            "kpi": kpi,
            "sql": sql_query,
            "results": results_text
        })
        print(f"  ✓ Done — {len(results)} rows")

    except Exception as e:
        conn.rollback()
        print(f"  ✗ Failed: {e}")
        kpi_results.append({
            "kpi": kpi,
            "sql": sql_query,
            "results": f"Query failed: {e}"
        })

conn.close()
print("\nAll KPIs processed!")

# Build summary for Claude
kpi_summary = ""
for item in kpi_results:
    kpi_summary += f"\n## {item['kpi']}\n"
    kpi_summary += f"{item['results']}\n"

# Generate executive summary for email body
print("\nGenerating executive summary...")
summary_message = client.messages.create(
    model="claude-sonnet-4-6",
    max_tokens=512,
    messages=[
        {"role": "user", "content": f"""
        Based on this KPI data, write a brief executive summary for an email.
        Maximum 5 bullet points with the most important insights.
        Keep it concise — this will be read on mobile.
        No markdown tables, no SQL, just clear business insights.
        
        KPI Data:
        {kpi_summary}
        """}
    ]
)

# Generate full report for .docx
print("Generating full report...")
report_message = client.messages.create(
    model="claude-sonnet-4-6",
    max_tokens=2048,
    messages=[
        {"role": "user", "content": f"""
        You are a business analyst writing a professional report.
        Based on the following KPI data, write a detailed report.
        Use headings for each KPI section.
        Use bullet points for key findings.
        Add a brief executive summary at the top.
        Do NOT use markdown tables — use numbered lists or bullet points instead.
        Keep the tone professional and business-focused.
        
        KPI Data:
        {kpi_summary}
        """}
    ]
)

executive_summary = summary_message.content[0].text
report_text = report_message.content[0].text

print("Executive summary and full report generated!")

# Save report as .docx
print("\nCreating .docx report...")
doc = Document()
doc.add_heading("Automated Business Report", level=1)
doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
doc.add_paragraph(f"Database: {os.getenv('DB_NAME')}")

html_content = markdown.markdown(report_text)
tmp_bytes = html2docx(html_content, title="")
tmp_doc = Document(io.BytesIO(tmp_bytes.getvalue()))
for element in tmp_doc.element.body:
    doc.element.body.append(element)

report_filename = f"report_{datetime.now().strftime('%Y-%m-%d')}.docx"
doc.save(report_filename)
print(f"Report saved as {report_filename}")

# Send email
print("\nSending email...")
sender = os.getenv("EMAIL_SENDER")
recipient = os.getenv("EMAIL_RECIPIENT")
password = os.getenv("EMAIL_PASSWORD")

msg = MIMEMultipart()
msg["From"] = sender
msg["To"] = recipient
msg["Subject"] = f"Automated Business Report — {datetime.now().strftime('%d/%m/%Y')}"

# Convert markdown to HTML for email
html_email = markdown.markdown(executive_summary)
msg.attach(MIMEText(html_email, "html"))

with open(report_filename, "rb") as f:
    attachment = MIMEBase("application", "octet-stream")
    attachment.set_payload(f.read())
    encoders.encode_base64(attachment)
    attachment.add_header("Content-Disposition", f"attachment; filename={report_filename}")
    msg.attach(attachment)

with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
    server.login(sender, password)
    server.sendmail(sender, recipient, msg.as_string())

print("Email sent successfully!")